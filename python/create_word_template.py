# Word Template Creation Script
# This script creates a basic Word template for the exemption petition

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def create_exemption_template():
    """Creates a Word template for the exemption petition"""
    
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜM BAŞKANLIĞINA")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()  # Empty line
    
    # Introduction paragraphs
    intro1 = doc.add_paragraph()
    intro1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro1.add_run(
        "Sivas Cumhuriyet Üniversitesi Ön lisans ve Lisans Eğitim-Öğretim ve Sınav Yönetmeliği gereği, "
        "aşağıdaki çizelgede belirtilen dersleri, Bölümünüze kayıt yaptırmadan önce okuduğum yükseköğretim programında "
        "aldım ve başardım. Bölümünüz 8 yarıyıllık ders programında yer alan ve daha önce aldığım bu derslere eşdeğer "
        "olduğunu düşündüğüm derslerden (Çizelge 1) muaf olmak istiyorum. Daha önceki öğrenimimde aldığım ders ve "
        "içerikleri ektedir."
    )
    
    intro2 = doc.add_paragraph()
    intro2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro2.add_run(
        "Muafiyet işlemim ilgili kurul kararları ile tamamlandıktan sonra; yönetmelik gereği, her yarıyıl için muaf "
        "olduğum ders kredisi kadar, danışmanımın önerisi doğrultusunda, bir üst sınıftan (Çizelge 2) ders almak ve "
        "\"Sivas Cumhuriyet Üniversitesi Muafiyet ve İntibak İşlemleri Yönergesi/Uygulama Esasları\" gereği ilgili "
        "yıla/yarıyıla intibakımın yapılmasını istiyorum. Gereğini arz ederim."
    )
    
    doc.add_paragraph()  # Empty line
    
    # Çizelge 1 title
    title1 = doc.add_paragraph()
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title1.add_run("Çizelge 1")
    run.bold = True
    
    # Table 1 - Main exemption table
    table1 = doc.add_table(rows=3, cols=10)
    table1.style = 'Table Grid'
    
    # Table 1 headers
    headers1 = table1.rows[0].cells
    headers1[0].merge(headers1[4])  # Merge first 5 cells
    headers1[0].text = "MUAFİYET İSTENEN DERSİN ALINDIĞI"
    headers1[5].merge(headers1[9])  # Merge last 5 cells  
    headers1[5].text = "MUAF OLMAK İSTENEN DERS (Bilgisayar Müh. Böl. Dersi)"
    
    # Second row headers
    headers2 = table1.rows[1].cells
    headers2[0].text = "Ders Kodu"
    headers2[1].text = "Ders Adı"
    headers2[2].text = "T"
    headers2[3].text = "U"
    headers2[4].text = "K"
    headers2[5].text = "Ders Kodu"
    headers2[6].text = "Ders Adı"
    headers2[7].text = "T"
    headers2[8].text = "U"
    headers2[9].text = "K"
    
    # Add 15 empty rows for courses
    for i in range(15):
        table1.add_row()
    
    doc.add_page_break()
    
    # Çizelge 1 devam
    title1_cont = doc.add_paragraph()
    title1_cont.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title1_cont.add_run("Çizelge 1 (devam)")
    run.bold = True
    
    # Table 1 continued
    table1_cont = doc.add_table(rows=15, cols=10)
    table1_cont.style = 'Table Grid'
    
    doc.add_paragraph()  # Empty line
    
    # Çizelge 2 title
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("Çizelge 2")
    run.bold = True
    
    # Table 2 - Upper level courses
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    
    # Table 2 headers
    headers2 = table2.rows[0].cells
    headers2[0].text = "Yıl/Yarıyıl"
    headers2[1].text = "Ders Kodu"
    headers2[2].text = "Ders Adı"
    headers2[3].text = "Kredi (T-U-K)"
    headers2[4].text = "Danışman Onayı"
    
    # Add 10 empty rows
    for i in range(10):
        table2.add_row()
    
    doc.add_paragraph()  # Empty lines
    doc.add_paragraph()
    
    # Signature section
    sig_para1 = doc.add_paragraph()
    sig_para1.add_run("Tarih              :")
    
    sig_para2 = doc.add_paragraph()
    sig_para2.add_run("Öğrenci Adı Soyadı :")
    
    sig_para3 = doc.add_paragraph()
    sig_para3.add_run("Öğrenci Numarası   :")
    
    sig_para4 = doc.add_paragraph()
    sig_para4.add_run("İmza               : ____________________")
    
    doc.add_paragraph()
    
    # Intibak section
    intibak_para = doc.add_paragraph()
    intibak_para.add_run("Bir üst yıl/yarıyıla intibakımın yapılmasını istemiyorum")
    
    doc.add_paragraph()
    
    # Footer notes
    ek_para = doc.add_paragraph()
    ek_para.add_run("Ek: Not çizelgesi (transkript) ve onaylı ders içerikleri")
    ek_para.runs[0].font.size = Pt(10)
    
    notes = [
        "SCÜ Muafiyet ve İntibak İşlemleri Yönergesi Madde 6 ve Sivas Cumhuriyet Üniversitesi Ön Lisans Ve "
        "Lisans Eğitim-Öğretim ve Sınav Yönetmeliği uyarınca öğrenci, muaf olduğu ders kredisi kadar, bir üst sınıftan "
        "danışmanının uygun bulduğu dersi/dersleri alabilir. Bu nedenle, muafiyet talebiniz ile birlikte üst sınıftan "
        "(yarıyıldan) almayı planladığınız dersleri mutlaka danışmanınızın onayı ile belirleyiniz (Çizelge 2'yi "
        "onaylatınız).",
        
        "Daha önceki eğitim programında alınan bir dersten muaf olunabilmesi için dersin başarılmış olması gerekir.",
        
        "\"SCÜ Muafiyet ve İntibak İşlemleri Yönergesi\" gereği muafiyet işlemlerinin sonucuna göre, ilgili "
        "yıla/yarıyıla intibak yapılır.",
        
        "Muafiyet/intibak talebi, ilgili yönetim kurulu tarafından karara bağlanıncaya kadar, öğrenciler, "
        "muafiyet talebinde bulunduğu derse/derslere ve üst yılda/yarıyılda almayı planladığı (danışmanının uygun "
        "bulduğu) derslere devam ederler."
    ]
    
    for note in notes:
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = note_para.add_run(note)
        run.font.size = Pt(10)
    
    # Save the template
    doc.save("Muafiyet_dilekcesi.docx")
    print("✅ Word template created: Muafiyet_dilekcesi.docx")

if __name__ == "__main__":
    create_exemption_template()
