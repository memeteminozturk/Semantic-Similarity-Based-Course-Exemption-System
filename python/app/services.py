# app/services.py - Simple version without complex type annotations
from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
import tempfile
import logging

import numpy as np
from sentence_transformers import SentenceTransformer

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.repository import CourseRepository

__all__ = ["SimilarityService", "ExemptionWordBuilder", "WordGenerationService", "PdfGenerationService"]

###############################################################################
# SimilarityService (unchanged)
###############################################################################

class SimilarityService:
    """Utility wrapper around a SentenceTransformer for course comparison."""

    def __init__(self, model_name, threshold, repo):
        self.model = SentenceTransformer(model_name)
        self.threshold = threshold
        self._repo = repo

        if not self._repo._cache:
            raise ValueError("CourseRepository cache boş. 'await repo.load()' çalıştırılmadı.")
        self._int_codes = []
        self._int_embs = None
        self._build_internal_cache()

    def _build_internal_cache(self):
        self._int_codes = list(self._repo._cache.keys())
        contents = list(self._repo._cache.values())
        if not contents:
            raise ValueError("Dahili ders içeriği bulunamadı.")
        self._int_embs = self.model.encode(contents, convert_to_numpy=True)

    def auto_match(self, ext_texts):
        if self._int_embs.size == 0:
            raise RuntimeError("Internal cache oluşturulmadı.")
        ext_embs = self.model.encode(list(ext_texts), convert_to_numpy=True)
        sims = (ext_embs @ self._int_embs.T) / (
            np.linalg.norm(ext_embs, axis=1, keepdims=True) * np.linalg.norm(self._int_embs, axis=1)
        )
        return sims

    @staticmethod
    def _cos_sim(a, b):
        denom = np.linalg.norm(a) * np.linalg.norm(b)
        return 0.0 if denom == 0 else float(np.dot(a, b) / denom)

    def bulk_similarity(self, pairs):
        flat = [t for pair in pairs for t in pair]
        embeds = self.model.encode(flat, convert_to_numpy=True)
        results = []
        thr_pct = self.threshold * 100
        for i in range(0, len(embeds), 2):
            sim = self._cos_sim(embeds[i], embeds[i + 1])
            pct = round(sim * 100, 2)
            results.append((sim, pct, pct >= thr_pct))
        return results

###############################################################################
# ExemptionWordBuilder
###############################################################################

class ExemptionWordBuilder:
    """Word belgesi oluşturmak için yardımcı sınıf"""

    @staticmethod
    def build(personal_info, exemption_courses, upper_courses=None, generation_dt=None, template_path="Muafiyet_dilekcesi.docx"):
        # Verileri Word formatına dönüştür
        data = ExemptionWordBuilder.map_data(personal_info, exemption_courses, upper_courses, generation_dt)
        # Belgeyi oluştur ve doldur
        return ExemptionWordBuilder.fill_template(template_path, data)

    @staticmethod
    def map_data(personal_info, exemption_courses, upper_courses, generation_dt):
        # Önceki dersler ve muaf dersler için ayrı listeler oluştur
        onceki_dersler = []
        muaf_dersler = []
        
        for course in exemption_courses:
            # Kredi bilgisini parçala (T-U-K formatı)
            ext_credit = course.get("ext_credit", "0-0-0")
            if isinstance(ext_credit, (int, float)):
                ext_credit = f"{ext_credit}-0-{ext_credit}"
            
            int_credit = course.get("int_credit", "0-0-0")
            if isinstance(int_credit, (int, float)):
                int_credit = f"{int_credit}-0-{int_credit}"
            
            # Ders bilgilerini ekle
            onceki_dersler.append({
                "kod": course.get("ext_code", ""),
                "ad": course.get("ext_name", ""),
                "T": ext_credit.split('-')[0] if '-' in ext_credit else ext_credit,
                "U": ext_credit.split('-')[1] if '-' in ext_credit else "0",
                "K": ext_credit.split('-')[2] if '-' in ext_credit else ext_credit,
            })            
            muaf_dersler.append({
                "kod": course.get("int_code", ""),
                "ad": course.get("int_name", ""),
                "T": int_credit.split('-')[0] if '-' in int_credit else int_credit,
                "U": int_credit.split('-')[1] if '-' in int_credit else "0",
                "K": int_credit.split('-')[2] if '-' in int_credit else int_credit,
            })
        
        return {
            "onceki_dersler": onceki_dersler,
            "muaf_dersler": muaf_dersler,
            "tarih": (generation_dt or datetime.now()).strftime("%d/%m/%Y"),
            "ad_soyad": f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}",
            "ogrenci_no": personal_info.get('student_number', ''),
            "universite": personal_info.get('university', ''),
            "fakulte": personal_info.get('faculty', ''),
            "bolum": personal_info.get('department', ''),
            "intibak_istiyor": True
        }

    @staticmethod
    def fill_template(template_path, data):
        doc = Document(template_path)
        tables = doc.tables
        
        # Çizelge 1'i doldur (ana tablo - Tablo 0)
        if len(tables) > 0:
            ExemptionWordBuilder.fill_table1(tables[0], data)
        
        # Çizelge 1 devamını doldur (ikinci tablo - Tablo 1)
        if len(tables) > 1:
            ExemptionWordBuilder.fill_table1_continued(tables[1], data)
        
        # Metin alanlarını doldur
        for p in doc.paragraphs:
            text = p.text.strip()
            if "Tarih" in text and ":" in text:
                p.clear()
                p.add_run(f"Tarih                          : {data['tarih']}")
            elif "Öğrenci Adı Soyadı" in text and ":" in text:
                p.clear()
                p.add_run(f"Öğrenci Adı Soyadı   : {data['ad_soyad']}")
            elif "Öğrenci Numarası" in text and ":" in text:
                p.clear()
                p.add_run(f"Öğrenci Numarası     : {data['ogrenci_no']}")
          # Bellekte sakla ve döndür
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
        
    @staticmethod
    def fill_table1(table, data):
        # Fill university, faculty, and department information in Row 1 (columns 0-4)
        # These cells are merged but we need to set them individually
        if len(table.rows) > 1:
            row1 = table.rows[1]
            
            # Clear existing content in the cells first
            for col_idx in range(5):  
                if col_idx < len(row1.cells):
                    row1.cells[col_idx].text = ""
            
            # Add formatted paragraphs to the first cell
            # This will ensure proper formatting without excessive tabs/spacing
            if len(row1.cells) > 0:
                cell = row1.cells[0]
                
                # Add each field as a separate paragraph with proper formatting
                uni_para = cell.paragraphs[0]
                uni_para.text = f"Üniversite: {data.get('universite', '')}"
                uni_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                fac_para = cell.add_paragraph()
                fac_para.text = f"Fakülte/MY/YO: {data.get('fakulte', '')}"
                fac_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                dept_para = cell.add_paragraph()
                dept_para.text = f"Bölüm: {data.get('bolum', '')}"
                dept_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Apply consistent formatting to all paragraphs
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)  # Set consistent font size
        
        # Fill course rows
        start_row = 3
        max_rows_on_first_table = 14
        courses_in_table = []
        
        for i, (prev, exempt) in enumerate(zip(data["onceki_dersler"], data["muaf_dersler"])):
            if i >= max_rows_on_first_table:
                break
                
            row_idx = start_row + i
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                ExemptionWordBuilder.fill_row(row, prev, exempt)
                courses_in_table.append((prev, exempt))

        # Calculate and fill totals
        if courses_in_table:
            # Calculate external totals (MUAFİYET İSTENEN)
            total_ext_T = sum(float(prev['T']) for prev, _ in courses_in_table)
            total_ext_U = sum(float(prev['U']) for prev, _ in courses_in_table)
            total_ext_K = sum(float(prev['K']) for prev, _ in courses_in_table)
            
            # Calculate internal totals (MUAF OLUNAN)
            total_int_T = sum(float(exempt['T']) for _, exempt in courses_in_table)
            total_int_U = sum(float(exempt['U']) for _, exempt in courses_in_table)
            total_int_K = sum(float(exempt['K']) for _, exempt in courses_in_table)
            
            # Find the totals row (last row)
            totals_row = table.rows[-1] if table.rows else None
            
            if totals_row and len(totals_row.cells) >= 10:
                # External totals (MUAFİYET İSTENEN)
                totals_row.cells[2].text = str(int(total_ext_T))  # T
                totals_row.cells[3].text = str(int(total_ext_U))  # U
                totals_row.cells[4].text = str(int(total_ext_K))  # K
                
                # Internal totals (MUAF OLUNAN)
                totals_row.cells[7].text = str(int(total_int_T))  # T
                totals_row.cells[8].text = str(int(total_int_U))  # U
                totals_row.cells[9].text = str(int(total_int_K))  # K

    @staticmethod
    def fill_table1_continued(table, data):
        max_rows_on_first_table = 14
        start_index = max_rows_on_first_table
        start_row = 1
        courses_in_table = []
        
        courses_remaining = data["onceki_dersler"][start_index:]
        exempt_remaining = data["muaf_dersler"][start_index:]
        
        for i, (prev, exempt) in enumerate(zip(courses_remaining, exempt_remaining)):
            row_idx = start_row + i
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                ExemptionWordBuilder.fill_row(row, prev, exempt)
                courses_in_table.append((prev, exempt))

        # Calculate and fill totals
        if courses_in_table:
            # Calculate external totals (MUAFİYET İSTENEN)
            total_ext_T = sum(float(prev['T']) for prev, _ in courses_in_table)
            total_ext_U = sum(float(prev['U']) for prev, _ in courses_in_table)
            total_ext_K = sum(float(prev['K']) for prev, _ in courses_in_table)
            
            # Calculate internal totals (MUAF OLUNAN)
            total_int_T = sum(float(exempt['T']) for _, exempt in courses_in_table)
            total_int_U = sum(float(exempt['U']) for _, exempt in courses_in_table)
            total_int_K = sum(float(exempt['K']) for _, exempt in courses_in_table)
            
            # Find the totals row (last row)
            totals_row = table.rows[-1] if table.rows else None
            
            if totals_row and len(totals_row.cells) >= 10:
                # External totals (MUAFİYET İSTENEN)
                totals_row.cells[2].text = str(int(total_ext_T))  # T
                totals_row.cells[3].text = str(int(total_ext_U))  # U
                totals_row.cells[4].text = str(int(total_ext_K))  # K
                
                # Internal totals (MUAF OLUNAN)
                totals_row.cells[7].text = str(int(total_int_T))  # T
                totals_row.cells[8].text = str(int(total_int_U))  # U
                totals_row.cells[9].text = str(int(total_int_K))  # K

    @staticmethod
    def fill_row(row, prev_course, exempt_course):
        cells = row.cells
        
        # Önceki okul dersi (0-4 sütunlar)
        cells[0].text = prev_course["kod"]
        cells[1].text = prev_course["ad"]
        cells[2].text = str(prev_course["T"])
        cells[3].text = str(prev_course["U"])
        cells[4].text = str(prev_course["K"])
        
        # SCÜ Bilgisayar Müh. dersi (5-9 sütunlar)
        cells[5].text = exempt_course["kod"]
        cells[6].text = exempt_course["ad"]
        cells[7].text = str(exempt_course["T"])
        cells[8].text = str(exempt_course["U"])
        cells[9].text = str(exempt_course["K"])

###############################################################################
# WordGenerationService
###############################################################################

class WordGenerationService:
    """Eski API ile uyumlu Word belgesi oluşturma servisi"""

    @staticmethod
    def generate_exemption_word(personal_info, muafiyet_courses, upper_courses=None, timestamp=None, template_path="Muafiyet_dilekcesi.docx"):
        # Tarih dönüşümü
        generation_dt = None
        if timestamp:
            if isinstance(timestamp, str):
                try:
                    generation_dt = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
                except ValueError:
                    generation_dt = datetime.now()
            elif isinstance(timestamp, datetime):
                generation_dt = timestamp
          # Alan adlarını eşle
        mapped_personal_info = {
            "first_name": personal_info.get("firstName", ""),
            "last_name": personal_info.get("lastName", ""),
            "student_number": personal_info.get("studentNumber", ""),
            "university": personal_info.get("university", ""),
            "faculty": personal_info.get("faculty", ""),
            "department": personal_info.get("department", ""),
        }
        
        # Ders verilerini eşle
        mapped_exemption_courses = []
        for course in muafiyet_courses:
            mapped_course = {
                "ext_university": course.get("ext_university", ""),
                "ext_code": course.get("ext_code", ""),
                "ext_name": course.get("ext_name", ""),
                "ext_credit": course.get("ext_credit", 0),
                "int_code": course.get("int_code", ""),
                "int_name": course.get("int_name", ""),
                "int_credit": course.get("int_credit", 0),
            }
            mapped_exemption_courses.append(mapped_course)
        
        # Word belgesini oluştur
        return ExemptionWordBuilder.build(
            personal_info=mapped_personal_info,
            exemption_courses=mapped_exemption_courses,
            upper_courses=None,
            generation_dt=generation_dt,
            template_path=template_path,
        )

###############################################################################
# PdfGenerationService (Backward compatibility)
###############################################################################

class PdfGenerationService:
    """Generates Word documents with backward compatibility."""

    @staticmethod
    def generate_exemption_pdf(personal_info, muafiyet_courses, upper_courses=None, timestamp=None):
        return WordGenerationService.generate_exemption_word(
            personal_info=personal_info,
            muafiyet_courses=muafiyet_courses,
            upper_courses=upper_courses,
            timestamp=timestamp,
        )


